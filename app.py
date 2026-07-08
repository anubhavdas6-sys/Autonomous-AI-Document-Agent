import os
import streamlit as st
from typing import List
from pydantic import BaseModel, Field
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- 1. FORCE THE API KEY CONFIGURATION ---
# Replace with your actual Gemini API key
if "GEMINI_API_KEY" not in os.environ:
    if "GEMINI_API_KEY" in st.secrets:
        os.environ["GEMINI_API_KEY"] = st.secrets["GEMINI_API_KEY"]

# Page Configuration
st.set_page_config(page_title="Autonomous AI Agent", page_icon="🤖", layout="centered")

# --- 2. AGENT STRUCTURAL SCHEMAS ---
class DocSection(BaseModel):
    title: str = Field(description="The heading/title of this document section.")
    content: str = Field(description="The comprehensive text, paragraphs, or bullet points for this section.")

class ExecutionPlan(BaseModel):
    document_title: str = Field(description="The overall title of the business document.")
    document_type: str = Field(description="Type of document (e.g., Proposal, SOP, Project Plan).")
    tasks_todo: List[str] = Field(description="The internal list of sub-tasks the agent needs to execute.")
    sections: List[DocSection] = Field(description="The actual generated sections of the document.")

# --- 3. DOCUMENT STYLING HELPER ---
def create_polished_document(plan: ExecutionPlan, filename: str) -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    COLOR_PRIMARY = RGBColor(0, 32, 96)   # Deep Navy
    COLOR_SECONDARY = RGBColor(127, 127, 127) # Muted Gray

    # Document Header/Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(plan.document_title.upper())
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f"Generated Autonomous {plan.document_type}")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_SECONDARY
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Content Sections
    for sec in plan.sections:
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        
        h_run = h.add_run(sec.title)
        h_run.font.name = 'Arial'
        h_run.font.size = Pt(14)
        h_run.font.bold = True
        h_run.font.color.rgb = COLOR_PRIMARY
        
        lines = sec.content.split('\n')
        for line in lines:
            if not line.strip():
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            
            if line.strip().startswith(('*', '-', '•')):
                p.paragraph_format.left_indent = Inches(0.25)
                run = p.add_run(line.strip().lstrip('*-• ').strip())
            else:
                run = p.add_run(line.strip())
                
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    doc.save(filename)
    return filename

# --- 4. STREAMLIT FRONTEND USER INTERFACE ---
st.title("🤖 Autonomous AI Document Agent")
st.write("Provide a description of the document you need. The agent will formulate a plan, execute the required tasks, and generate a formatted Word document.")

# Input text area
user_request = st.text_area(
    "What document would you like to create?",
    placeholder="e.g., Create a comprehensive project proposal for a new mobile application targeting local grocery delivery, focusing on architecture, budget, and a 6-month roadmap.",
    height=120
)

if st.button("Execute Agent Workflow", type="primary"):
    if not user_request.strip():
        st.error("Please enter a request before running the agent!")
    else:
        with st.spinner("🤖 Agent is thinking, constructing execution plan, and running tasks..."):
            try:
                # 1. Initialize the LLM
                llm = ChatGoogleGenerativeAI(model="gemini-2.5-flash", temperature=0.1)
                
                # 2. Force the model to natively output our Pydantic structural schema
                structured_llm = llm.with_structured_output(ExecutionPlan)
                
                # 3. Simplify the prompt since the schema constraint is enforced natively
                prompt = ChatPromptTemplate.from_messages([
                    ("system", (
                        "You are an autonomous business operations agent. Your job is to take a raw user request, "
                        "determine what kind of document needs to be created, outline a Todo task list of actions "
                        "required to fulfill it, execute those actions internally, and output a highly comprehensive, "
                        "professional business report layout matching the requested fields."
                    )),
                    ("user", "Fulfill this request comprehensively: {user_request}")
                ])
                
                # Create the clean chain
                chain = prompt | structured_llm
                
                # 4. Invoke the chain (this returns a clean, fully-formed ExecutionPlan object directly!)
                plan = chain.invoke({"user_request": user_request})
                
                if not plan:
                    raise ValueError("The AI model returned an empty response. Please try tweaking your prompt words.")

                # Success Display
                st.success("✅ Workflow Executed Successfully!")
                
                # Display structural details from Agent Logs
                st.subheader("📋 Agent Internal Execution Logs")
                st.markdown(f"**Document Type Determined:** `{plan.document_type}`")
                st.markdown(f"**Document Title:** *{plan.document_title}*")
                
                st.write("**Executed Task Checklist:**")
                for task in plan.tasks_todo:
                    st.markdown(f"- ✅ {task}")
                
                # Compile document
                output_filename = "Streamlit_Agent_Output.docx"
                create_polished_document(plan, filename=output_filename)
                
                # Provide direct download button within app
                with open(output_filename, "rb") as file:
                    st.download_button(
                        label="📥 Download Polished Document (.docx)",
                        data=file,
                        file_name=f"{plan.document_title.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
            except Exception as e:
                st.error(f"An error occurred during runtime: {e}")